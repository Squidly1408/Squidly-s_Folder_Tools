import os
from tkinter import (
    Tk,
    filedialog,
    messagebox,
    Toplevel,
    Label,
    Button,
    Frame,
    StringVar,
    OptionMenu,
)
from fpdf import FPDF
from PyPDF2 import PdfReader
from PIL import Image
from pydub import AudioSegment
from docx import Document
import openpyxl
import pandas as pd
from docx2pdf import convert


def themed_error_message(error_message):
    """Display error message in a themed popup window."""
    error_window = Toplevel(root)
    error_window.title("Error")
    error_window.geometry("400x200")

    # Customize the appearance
    error_frame = Frame(error_window, bg="#f2f2f2")  # Set a light background color
    error_frame.pack(fill="both", expand=True)

    error_label = Label(
        error_frame,
        text=error_message,
        bg="#f2f2f2",
        font=("Arial", 12),
        wraplength=380,
    )
    error_label.pack(pady=20)

    ok_button = Button(
        error_frame,
        text="OK",
        command=error_window.destroy,
        bg="#007BFF",
        fg="white",
        font=("Arial", 10),
    )
    ok_button.pack(pady=10)


def convert_files(source_files, destination_format):
    # Determine the source format from the first file's extension
    if not source_files:
        themed_error_message("No source files provided.")
        return

    source_format = os.path.splitext(source_files[0])[1][
        1:
    ]  # Get the file extension without the dot
    converters = {
        "txt": {
            "pdf": convert_txt_to_pdf,
            "docx": convert_txt_to_docx,
            "xlsx": convert_txt_to_xlsx,
        },
        "pdf": {"txt": convert_pdf_to_txt, "docx": convert_pdf_to_docx},
        "docx": {"pdf": convert_docx_to_pdf, "txt": convert_docx_to_txt},
        "jpg": {"png": convert_jpg_to_png, "gif": convert_jpg_to_gif},
        "png": {"jpg": convert_png_to_jpg, "gif": convert_png_to_gif},
        "wav": {"mp3": convert_wav_to_mp3, "ogg": convert_wav_to_ogg},
        "mp3": {"wav": convert_mp3_to_wav, "ogg": convert_mp3_to_ogg},
        "xlsx": {"csv": convert_xlsx_to_csv},
    }

    if (
        source_format not in converters
        or destination_format not in converters[source_format]
    ):
        themed_error_message(
            f"Conversion from {source_format} to {destination_format} is not supported."
        )
        return

    for source_file in source_files:
        destination_file = os.path.splitext(source_file)[0] + f".{destination_format}"
        try:
            converters[source_format][destination_format](source_file, destination_file)
        except Exception as e:
            themed_error_message(
                f"Error converting {source_file} to {destination_format}: {str(e)}"
            )

    messagebox.showinfo(
        "Conversion complete", "All selected files have been converted."
    )


# Conversion functions
def convert_txt_to_pdf(source_file, destination_file):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    with open(source_file, "r") as txt_file:
        for line in txt_file:
            pdf.cell(200, 10, txt=line, ln=True)
    pdf.output(destination_file)


def convert_pdf_to_txt(source_file, destination_file):
    reader = PdfReader(source_file)
    with open(destination_file, "w") as txt_file:
        for page in reader.pages:
            txt_file.write(page.extract_text())


def convert_pdf_to_docx(source_file, destination_file):
    convert(source_file, destination_file)


def convert_docx_to_pdf(source_file, destination_file):
    convert(source_file, destination_file)


def convert_docx_to_txt(source_file, destination_file):
    doc = Document(source_file)
    with open(destination_file, "w") as txt_file:
        for para in doc.paragraphs:
            txt_file.write(para.text + "\n")


def convert_txt_to_docx(source_file, destination_file):
    doc = Document()
    with open(source_file, "r") as txt_file:
        for line in txt_file:
            doc.add_paragraph(line)
    doc.save(destination_file)


def convert_txt_to_xlsx(source_file, destination_file):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    with open(source_file, "r") as txt_file:
        for row_index, line in enumerate(txt_file, start=1):
            sheet.cell(row=row_index, column=1, value=line.strip())
    workbook.save(destination_file)


def convert_jpg_to_png(source_file, destination_file):
    img = Image.open(source_file)
    img.save(destination_file, "PNG")


def convert_jpg_to_gif(source_file, destination_file):
    img = Image.open(source_file)
    img.save(destination_file, "GIF")


def convert_png_to_jpg(source_file, destination_file):
    img = Image.open(source_file)
    img.convert("RGB").save(destination_file, "JPEG")


def convert_png_to_gif(source_file, destination_file):
    img = Image.open(source_file)
    img.save(destination_file, "GIF")


def convert_wav_to_mp3(source_file, destination_file):
    sound = AudioSegment.from_wav(source_file)
    sound.export(destination_file, format="mp3")


def convert_wav_to_ogg(source_file, destination_file):
    sound = AudioSegment.from_wav(source_file)
    sound.export(destination_file, format="ogg")


def convert_mp3_to_wav(source_file, destination_file):
    sound = AudioSegment.from_mp3(source_file)
    sound.export(destination_file, format="wav")


def convert_mp3_to_ogg(source_file, destination_file):
    sound = AudioSegment.from_mp3(source_file)
    sound.export(destination_file, format="ogg")


def convert_xlsx_to_csv(source_file, destination_file):
    df = pd.read_excel(source_file)
    df.to_csv(destination_file, index=False)


def select_files():
    """Open file dialog to select files and choose the destination format."""
    source_files = filedialog.askopenfilenames(title="Select Files")
    if not source_files:
        return

    # Allow user to select destination format
    format_var = StringVar()
    format_var.set("pdf")  # Set default format

    format_window = Toplevel(root)
    format_window.title("Select Destination Format")
    format_window.geometry("300x150")

    format_label = Label(
        format_window, text="Choose destination format:", font=("Arial", 12)
    )
    format_label.pack(pady=10)

    format_options = [
        "pdf",
        "txt",
        "docx",
        "xlsx",
        "png",
        "jpg",
        "gif",
        "mp3",
        "wav",
        "ogg",
        "csv",
    ]
    format_menu = OptionMenu(format_window, format_var, *format_options)
    format_menu.pack(pady=10)

    convert_button = Button(
        format_window,
        text="Convert",
        command=lambda: convert_files(source_files, format_var.get()),
    )
    convert_button.pack(pady=20)


# Entry point
if __name__ == "__main__":
    root = Tk()
    root.withdraw()  # Hide the root window as we will use dialogs only

    select_files()  # Start the file selection process

    root.mainloop()
